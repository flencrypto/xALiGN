"""Document parsing service.

Extracts plain text and structured requirements from PDF and Word (.docx) files.
Uses pdfplumber for PDFs and python-docx for Word documents.
"""

import io
import json
import logging
import re
from pathlib import Path
from typing import Any

logger = logging.getLogger("align.document_parser")

# ── Requirement extraction heuristics ─────────────────────────────────────────

_REQ_PATTERNS: list[tuple[re.Pattern, str]] = [
    (re.compile(r"\b(shall|must|is required to|are required to)\b", re.I), "mandatory"),
    (re.compile(r"\b(should|is expected to|are expected to)\b", re.I), "preferred"),
    (re.compile(r"\b(may|can|is permitted to)\b", re.I), "optional"),
    (re.compile(r"\b(will not|shall not|must not)\b", re.I), "exclusion"),
]

_MIN_SENTENCE_LEN = 20
_MAX_SENTENCE_LEN = 500


def _split_sentences(text: str) -> list[str]:
    """Split text into sentences on common delimiters."""
    raw = re.split(r"(?<=[.!?])\s+|\n{2,}", text)
    return [s.strip() for s in raw if _MIN_SENTENCE_LEN <= len(s.strip()) <= _MAX_SENTENCE_LEN]


def _classify_sentence(sentence: str) -> str | None:
    """Return compliance category if sentence looks like a requirement, else None."""
    for pattern, category in _REQ_PATTERNS:
        if pattern.search(sentence):
            return category
    return None


def extract_requirements(text: str, default_category: str = "general") -> list[dict[str, Any]]:
    """
    Extract structured requirements from plain text.

    Returns a list of dicts with keys: ``requirement``, ``category``.
    """
    requirements: list[dict[str, Any]] = []
    seen: set[str] = set()
    for sentence in _split_sentences(text):
        cat = _classify_sentence(sentence) or default_category
        norm = sentence.lower()
        if norm not in seen:
            seen.add(norm)
            requirements.append({"requirement": sentence, "category": cat})
    return requirements


# ── PDF parsing ───────────────────────────────────────────────────────────────

def parse_pdf(file_bytes: bytes) -> tuple[str, list[dict[str, Any]]]:
    """
    Parse a PDF file and return ``(content_text, requirements)``.

    Falls back to empty results if pdfplumber is unavailable or parsing fails.
    """
    try:
        import pdfplumber  # optional heavy dependency
    except ImportError:
        logger.warning("pdfplumber not installed; skipping PDF parsing")
        return "", []

    pages_text: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages_text.append(page_text)
    except Exception as exc:
        logger.error("PDF parsing failed: %s", exc)
        return "", []

    content_text = "\n\n".join(pages_text)
    requirements = extract_requirements(content_text)
    return content_text, requirements


# ── Word parsing ──────────────────────────────────────────────────────────────

def parse_docx(file_bytes: bytes) -> tuple[str, list[dict[str, Any]]]:
    """
    Parse a .docx file and return ``(content_text, requirements)``.

    Extracts text from paragraphs and table cells.
    Falls back to empty results if python-docx is unavailable or parsing fails.
    """
    try:
        from docx import Document  # optional heavy dependency
    except ImportError:
        logger.warning("python-docx not installed; skipping Word parsing")
        return "", []

    parts: list[str] = []
    try:
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
    except Exception as exc:
        logger.error("Word parsing failed: %s", exc)
        return "", []

    content_text = "\n\n".join(parts)
    requirements = extract_requirements(content_text)
    return content_text, requirements


# ── Dispatcher ────────────────────────────────────────────────────────────────

def parse_document(
    file_bytes: bytes,
    filename: str,
) -> tuple[str, str]:
    """
    Parse a document by file extension.

    Returns ``(content_text, extracted_requirements_json)`` where
    ``extracted_requirements_json`` is a JSON-serialised list of
    ``{"requirement": str, "category": str}`` dicts.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        content_text, reqs = parse_pdf(file_bytes)
    elif ext in {".docx", ".doc"}:
        content_text, reqs = parse_docx(file_bytes)
    else:
        logger.info("No parser for extension %s; returning empty extraction", ext)
        content_text, reqs = "", []

    return content_text, json.dumps(reqs)
