"""Export service – generates Pursuit Pack PDFs, Tender Response Word docs,
and Compliance Matrix Excel files.

Dependencies:
  - reportlab (PDF)
  - python-docx (Word)
  - openpyxl (Excel)
"""

import io
import logging
from datetime import datetime, timezone
from typing import Any

logger = logging.getLogger("align.export")


# ── Helpers ───────────────────────────────────────────────────────────────────

def _utcnow() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


# ── PDF – Pursuit Pack ────────────────────────────────────────────────────────

def build_pursuit_pack_pdf(opportunity: Any, bid: Any, compliance_items: list[Any]) -> bytes:
    """
    Generate a Pursuit Pack PDF for a bid opportunity.

    Includes: cover page, opportunity summary, bid details, compliance overview.
    Returns raw PDF bytes.
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            HRFlowable,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError:
        raise RuntimeError("reportlab is not installed; cannot generate PDF exports.")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    normal = styles["Normal"]

    accent = ParagraphStyle(
        "accent",
        parent=normal,
        textColor=colors.HexColor("#1E3A5F"),
        fontSize=9,
        leading=14,
    )

    elements = []

    # ── Cover ──
    elements.append(Spacer(1, 1 * cm))
    elements.append(Paragraph("PURSUIT PACK", h1))
    elements.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#1E3A5F")))
    elements.append(Spacer(1, 0.4 * cm))

    opp_name = getattr(opportunity, "name", "Unnamed Opportunity")
    bid_title = getattr(bid, "title", "Untitled Bid")
    elements.append(Paragraph(f"<b>Opportunity:</b> {opp_name}", normal))
    elements.append(Paragraph(f"<b>Bid Title:</b> {bid_title}", normal))
    tender_ref = getattr(bid, "tender_ref", None)
    if tender_ref:
        elements.append(Paragraph(f"<b>Tender Ref:</b> {tender_ref}", normal))
    bid_status = getattr(bid, "status", "")
    elements.append(Paragraph(f"<b>Status:</b> {bid_status}", normal))
    elements.append(Paragraph(f"<b>Generated:</b> {_utcnow()}", accent))
    elements.append(Spacer(1, 0.5 * cm))

    # ── Win Themes ──
    win_themes = getattr(bid, "win_themes", None)
    if win_themes:
        elements.append(Paragraph("Win Themes", h2))
        elements.append(Paragraph(win_themes, normal))
        elements.append(Spacer(1, 0.3 * cm))

    # ── Compliance Summary ──
    if compliance_items:
        elements.append(Paragraph("Compliance Summary", h2))
        table_data = [["#", "Requirement", "Status", "Category"]]
        for i, item in enumerate(compliance_items[:50], 1):
            req = getattr(item, "requirement", "")
            req_short = req[:80] + "…" if len(req) > 80 else req
            table_data.append([
                str(i),
                req_short,
                getattr(item, "compliance_status", ""),
                getattr(item, "category", "") or "",
            ])
        tbl = Table(table_data, colWidths=[1 * cm, 9 * cm, 3 * cm, 3 * cm])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A5F")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F0F4F8")]),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        elements.append(tbl)

    doc.build(elements)
    buf.seek(0)
    return buf.read()


# ── Word – Tender Response Pack ───────────────────────────────────────────────

def build_tender_response_pack_docx(bid: Any, compliance_items: list[Any], rfis: list[Any]) -> bytes:
    """
    Generate a Tender Response Pack Word document.

    Includes: bid cover, compliance answers, RFI log.
    Returns raw .docx bytes.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx is not installed; cannot generate Word exports.")

    doc = Document()

    # ── Title ──
    title = doc.add_heading("Tender Response Pack", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    bid_title = getattr(bid, "title", "Untitled Bid")
    tender_ref = getattr(bid, "tender_ref", "")
    doc.add_paragraph(f"Bid: {bid_title}" + (f"  |  Ref: {tender_ref}" if tender_ref else ""))
    doc.add_paragraph(f"Generated: {_utcnow()}")
    doc.add_paragraph("")

    # ── Compliance Matrix ──
    if compliance_items:
        doc.add_heading("Compliance Matrix", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Requirement"
        hdr[1].text = "Category"
        hdr[2].text = "Status"
        hdr[3].text = "Evidence / Answer"
        for item in compliance_items:
            row = table.add_row().cells
            row[0].text = getattr(item, "requirement", "")
            row[1].text = getattr(item, "category", "") or ""
            row[2].text = getattr(item, "compliance_status", "")
            row[3].text = getattr(item, "evidence", "") or ""
        doc.add_paragraph("")

    # ── RFI Log ──
    if rfis:
        doc.add_heading("RFI Log", level=1)
        rfi_table = doc.add_table(rows=1, cols=4)
        rfi_table.style = "Table Grid"
        hdr = rfi_table.rows[0].cells
        hdr[0].text = "Question"
        hdr[1].text = "Category"
        hdr[2].text = "Priority"
        hdr[3].text = "Answer"
        for rfi in rfis:
            row = rfi_table.add_row().cells
            row[0].text = getattr(rfi, "question", "")
            row[1].text = getattr(rfi, "category", "") or ""
            row[2].text = getattr(rfi, "priority", "")
            row[3].text = getattr(rfi, "answer", "") or ""

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Excel – Compliance Matrix ─────────────────────────────────────────────────

def build_compliance_matrix_xlsx(bid: Any, compliance_items: list[Any]) -> bytes:
    """
    Generate a Compliance Matrix Excel workbook.

    Sheet 1: Summary. Sheet 2: Full compliance items.
    Returns raw .xlsx bytes.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise RuntimeError("openpyxl is not installed; cannot generate Excel exports.")

    wb = Workbook()
    wb.remove(wb.active)  # remove default sheet

    # ── Summary sheet ──
    ws_sum = wb.create_sheet("Summary")
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="1E3A5F")

    ws_sum.append(["Tender Response – Compliance Matrix"])
    ws_sum["A1"].font = Font(bold=True, size=14)
    ws_sum.append(["Bid:", getattr(bid, "title", "")])
    tender_ref = getattr(bid, "tender_ref", "")
    ws_sum.append(["Tender Ref:", tender_ref or ""])
    ws_sum.append(["Generated:", _utcnow()])
    ws_sum.append([])

    total = len(compliance_items)
    by_status: dict[str, int] = {}
    for item in compliance_items:
        s = getattr(item, "compliance_status", "tbc")
        by_status[s] = by_status.get(s, 0) + 1

    ws_sum.append(["Status", "Count", "% of Total"])
    ws_sum["A6"].font = Font(bold=True)
    ws_sum["B6"].font = Font(bold=True)
    ws_sum["C6"].font = Font(bold=True)
    for s, count in sorted(by_status.items()):
        pct = round(count / total * 100, 1) if total else 0.0
        ws_sum.append([s, count, f"{pct}%"])

    # ── Full items sheet ──
    ws = wb.create_sheet("Compliance Items")
    headers = ["#", "Requirement", "Category", "Status", "Evidence", "Owner", "Notes"]
    ws.append(headers)
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    for i, item in enumerate(compliance_items, 1):
        ws.append([
            i,
            getattr(item, "requirement", ""),
            getattr(item, "category", "") or "",
            getattr(item, "compliance_status", ""),
            getattr(item, "evidence", "") or "",
            getattr(item, "owner", "") or "",
            getattr(item, "notes", "") or "",
        ])

    # Auto-fit columns
    col_widths = [6, 60, 20, 12, 40, 20, 30]
    for col_idx, width in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()
